"""上游 AI 解析路由 (/api/parse)。

实现静态（离线）采集新场景的解析层：
- POST /api/parse/requirement  接收需求文本，调用联通 AI 解析智能体
- POST /api/parse/stream        SSE 流式版本，前端可实时展示 AI 思考过程
- GET  /api/parse/{parse_id}    查询解析结果与缺失字段清单
- POST /api/parse/{parse_id}/supplement  用户补全缺失字段
"""
from __future__ import annotations

import datetime as _dt
import json
import os
from typing import Any, Dict, List, Optional

from fastapi import APIRouter, Depends, File, HTTPException, UploadFile
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field
from sqlalchemy.orm import Session

from backend.database import get_db
from backend.models import AgentInvocation, MissingField, ParseResult
from backend.services.upstream import (
    UnicomParseClient,
    UnicomParseConfig,
    UnicomParseError,
    extract_structured_result,
)


router = APIRouter(prefix="/api/parse", tags=["parse"])


# ------ Schemas ---------------------------------------------------------------


class ParseReq(BaseModel):
    requirement: str = Field(..., description="采集需求 / 接口规范文本")
    task_id: Optional[int] = Field(None, description="绑定的采集任务 ID，可空")
    user_name: str = Field("aiFind", description="调用方用户名，审计用")
    conversation_id: Optional[str] = None


class SupplementReq(BaseModel):
    values: Dict[str, Any] = Field(default_factory=dict, description="字段路径 -> 值")


# ------ 缺失字段计算 -----------------------------------------------------------

REQUIRED_FIELDS: List[Dict[str, str]] = [
    {"path": "source_info.source_code", "reason": "源端连接器编码"},
    {"path": "source_info.source_file_path", "reason": "源端 SFTP 目录"},
    {"path": "source_info.source_row_split_char", "reason": "源端文件行间分隔符"},
    {"path": "source_info.source_col_split_char", "reason": "源端文件字段分隔符"},
    {"path": "source_info.source_ext", "reason": "源端文件后缀"},
    {"path": "target_info.target_code", "reason": "目标端连接器编码"},
    {"path": "target_info.target_database_name", "reason": "目标端数据库名称"},
    {"path": "target_info.table_name", "reason": "目标表名"},
    {"path": "task_base_info.logic_dir", "reason": "逻辑设计器目录"},
    {"path": "task_base_info.logic_project", "reason": "逻辑设计器项目"},
    {"path": "task_base_info.logic_data_level", "reason": "数据分层"},
    {"path": "task_base_info.logic_topic_name", "reason": "主题域"},
    {"path": "task_base_info.logic_type", "reason": "业务分类"},
    {"path": "task_base_info.physical_dir", "reason": "物理设计器目录"},
]


def _lookup(obj: Dict[str, Any], path: str) -> Any:
    cur: Any = obj
    for part in path.split("."):
        if not isinstance(cur, dict):
            return None
        cur = cur.get(part)
    return cur


def compute_missing_fields(structured: Dict[str, Any]) -> List[Dict[str, str]]:
    missing: List[Dict[str, str]] = []
    for item in REQUIRED_FIELDS:
        v = _lookup(structured, item["path"])
        if v in (None, "", [], {}):
            missing.append({"field_path": item["path"], "reason": item["reason"]})
    ts = structured.get("table_structure")
    if not isinstance(ts, list) or not ts:
        missing.append({"field_path": "table_structure", "reason": "表结构字段列表"})
    return missing


# ------ 审计写入 --------------------------------------------------------------


def _audit(
    db: Session,
    *,
    task_id: Optional[int],
    tool_name: str,
    request: Dict[str, Any],
    response: Dict[str, Any],
    status: str,
    error: str = "",
    latency_ms: int = 0,
) -> None:
    db.add(
        AgentInvocation(
            task_id=task_id,
            direction="upstream",
            agent="unicom_ai_parser",
            tool_name=tool_name,
            request=_safe_json(request),
            response=_safe_json(response),
            status=status,
            error=error,
            latency_ms=latency_ms,
        )
    )


def _safe_json(obj: Any) -> Any:
    """把调用原始入出参落库前做一次 JSON 可序列化检查，避免污染 DB。"""
    try:
        json.dumps(obj, ensure_ascii=False, default=str)
        return obj
    except (TypeError, ValueError):
        return {"_repr": repr(obj)[:2000]}


def _client() -> UnicomParseClient:
    try:
        cfg = UnicomParseConfig.from_env()
    except UnicomParseError as e:
        raise HTTPException(status_code=500, detail=str(e)) from e
    return UnicomParseClient(cfg)


def _is_mock_mode() -> bool:
    return os.environ.get("UNICOM_MOCK", "").strip().lower() in {"1", "true", "yes"}


def _mock_result(requirement: str) -> Dict[str, Any]:
    """当配置缺失或显式开启 mock 时使用；覆盖典型字段用于本地调试。"""
    return {
        "source_info": {
            "source_code": "sftp_cb_224",
            "source_file_path": "/sftp/data/src_ai_caiji",
            "source_row_split_char": "UNIX换行符",
            "source_col_split_char": "逗号",
            "source_ext": "gz",
        },
        "target_info": {
            "target_code": "305实时数仓集群_hive",
            "target_database_name": "paimon_src",
            "table_name": "src_ai_caiji",
        },
        "ftp_info": {
            "host": "",
            "port": 22,
            "username": "",
            "file_path": "/sftp/data/src_ai_caiji",
        },
        "task_base_info": {
            "logic_dir": "治理监控测试",
            "logic_project": "经分",
            "logic_data_level": "SRC",
            "logic_topic_name": "CUS(客户域)",
            "logic_type": "5G",
            "physical_dir": "测试目录",
        },
        "table_structure": [
            {"name": "grid_id", "type": "text", "pk": True, "comment": "栅格ID"},
        ],
        "_source": "mock",
        "_echo": requirement[:200],
    }


# ------ 路由实现 --------------------------------------------------------------


@router.post("/requirement")
def parse_requirement(req: ParseReq, db: Session = Depends(get_db)):
    if not req.requirement.strip():
        raise HTTPException(status_code=400, detail="requirement 不能为空")

    started = _dt.datetime.utcnow()
    if _is_mock_mode():
        structured = _mock_result(req.requirement)
        content = json.dumps(structured, ensure_ascii=False)
        conv_id = req.conversation_id or "mock-conv"
        msg_id = "mock-msg"
        latency_ms = 0
        raw_response = json.dumps({"data": {"result": {"content": content}}}, ensure_ascii=False)
        pr = _persist_parse_result(
            db,
            req=req,
            content=content,
            structured=structured,
            raw_response=raw_response,
            conv_id=conv_id,
            msg_id=msg_id,
            latency_ms=latency_ms,
            status="succeeded",
        )
        _audit(
            db,
            task_id=req.task_id,
            tool_name="unicom_ai_parse_collection_requirement",
            request={"requirement": req.requirement, "mock": True},
            response={"structured": structured},
            status="succeeded",
            latency_ms=latency_ms,
        )
        db.commit()
        return _serialize(pr)

    client = _client()
    try:
        result = client.invoke(
            req.requirement,
            user_name=req.user_name,
            conversation_id=req.conversation_id,
        )
    except UnicomParseError as e:
        pr = _persist_parse_result(
            db,
            req=req,
            content="",
            structured={},
            raw_response="",
            conv_id=req.conversation_id or "",
            msg_id="",
            latency_ms=int((_dt.datetime.utcnow() - started).total_seconds() * 1000),
            status="failed",
            error=str(e),
        )
        _audit(
            db,
            task_id=req.task_id,
            tool_name="unicom_ai_parse_collection_requirement",
            request={"requirement": req.requirement},
            response={},
            status="failed",
            error=str(e),
        )
        db.commit()
        raise HTTPException(status_code=502, detail=f"上游调用失败: {e}") from e

    pr = _persist_parse_result(
        db,
        req=req,
        content=result.content,
        structured=result.structured,
        raw_response=result.raw_response,
        conv_id=result.conversation_id,
        msg_id=result.message_id,
        latency_ms=result.latency_ms,
        status="succeeded",
    )
    _audit(
        db,
        task_id=req.task_id,
        tool_name="unicom_ai_parse_collection_requirement",
        request={"requirement": req.requirement},
        response={
            "structured": result.structured,
            "conversation_id": result.conversation_id,
        },
        status="succeeded",
        latency_ms=result.latency_ms,
    )
    db.commit()
    return _serialize(pr)


@router.post("/stream")
def parse_stream(req: ParseReq, db: Session = Depends(get_db)):
    """SSE 流式解析。前端以 EventSource 或 fetch + ReadableStream 消费。"""
    if not req.requirement.strip():
        raise HTTPException(status_code=400, detail="requirement 不能为空")

    if _is_mock_mode():
        return StreamingResponse(
            _mock_stream(req, db),
            media_type="text/event-stream",
            headers={"Cache-Control": "no-cache"},
        )

    client = _client()

    def event_source() -> Any:
        try:
            final_payload: Dict[str, Any] = {}
            for evt in client.stream(
                req.requirement,
                user_name=req.user_name,
                conversation_id=req.conversation_id,
            ):
                if evt.get("type") == "final":
                    final_payload = evt
                yield _sse_line(evt)
            pr = _persist_parse_result(
                db,
                req=req,
                content=final_payload.get("content", ""),
                structured=final_payload.get("structured", {}),
                raw_response=final_payload.get("content", ""),
                conv_id=final_payload.get("conversation_id", ""),
                msg_id=final_payload.get("message_id", ""),
                latency_ms=final_payload.get("latency_ms", 0),
                status="succeeded",
            )
            _audit(
                db,
                task_id=req.task_id,
                tool_name="unicom_ai_parse_collection_requirement.stream",
                request={"requirement": req.requirement},
                response={"structured": pr.structured},
                status="succeeded",
                latency_ms=pr.latency_ms or 0,
            )
            db.commit()
            yield _sse_line({"type": "done", "parse_id": pr.id, "missing_fields": pr.missing_fields})
        except UnicomParseError as e:
            db.rollback()
            yield _sse_line({"type": "error", "message": str(e)})

    return StreamingResponse(event_source(), media_type="text/event-stream")


@router.get("/{parse_id}")
def get_parse_result(parse_id: int, db: Session = Depends(get_db)):
    pr = db.query(ParseResult).filter(ParseResult.id == parse_id).first()
    if not pr:
        raise HTTPException(404, "解析记录不存在")
    return _serialize(pr)


@router.post("/{parse_id}/supplement")
def supplement_parse_result(
    parse_id: int, body: SupplementReq, db: Session = Depends(get_db)
):
    pr = db.query(ParseResult).filter(ParseResult.id == parse_id).first()
    if not pr:
        raise HTTPException(404, "解析记录不存在")

    structured = dict(pr.structured or {})
    now = _dt.datetime.utcnow()
    existing = {
        row.field_path: row
        for row in db.query(MissingField)
        .filter(MissingField.parse_id == parse_id)
        .all()
    }
    for path, value in body.values.items():
        _set_by_path(structured, path, value)
        row = existing.get(path)
        if row is None:
            row = MissingField(
                parse_id=parse_id,
                task_id=pr.task_id,
                field_path=path,
                reason="用户补全",
            )
            db.add(row)
        row.user_value = json.dumps(value, ensure_ascii=False) if not isinstance(value, str) else value
        row.confirmed_at = now

    pr.structured = structured
    pr.missing_fields = compute_missing_fields(structured)
    db.commit()
    db.refresh(pr)
    return _serialize(pr)


# ------ Helpers ---------------------------------------------------------------


def _persist_parse_result(
    db: Session,
    *,
    req: ParseReq,
    content: str,
    structured: Dict[str, Any],
    raw_response: str,
    conv_id: str,
    msg_id: str,
    latency_ms: int,
    status: str,
    error: str = "",
) -> ParseResult:
    if not structured and content:
        structured = extract_structured_result(content)
    missing = compute_missing_fields(structured) if structured else [
        {"field_path": item["path"], "reason": item["reason"]} for item in REQUIRED_FIELDS
    ] + [{"field_path": "table_structure", "reason": "表结构字段列表"}]

    pr = ParseResult(
        task_id=req.task_id,
        source="unicom",
        requirement=req.requirement,
        content=content or "",
        structured=structured or {},
        missing_fields=missing,
        conversation_id=conv_id,
        message_id=msg_id,
        latency_ms=latency_ms,
        status=status,
        error=error,
        raw_response=(raw_response or "")[:200_000],
    )
    db.add(pr)
    db.flush()
    for item in missing:
        db.add(
            MissingField(
                parse_id=pr.id,
                task_id=req.task_id,
                field_path=item["field_path"],
                reason=item["reason"],
            )
        )
    return pr


def _serialize(pr: ParseResult) -> Dict[str, Any]:
    return {
        "id": pr.id,
        "task_id": pr.task_id,
        "source": pr.source,
        "status": pr.status,
        "error": pr.error,
        "requirement": pr.requirement,
        "content": pr.content,
        "structured": pr.structured or {},
        "missing_fields": pr.missing_fields or [],
        "conversation_id": pr.conversation_id,
        "message_id": pr.message_id,
        "latency_ms": pr.latency_ms,
        "created_at": str(pr.created_at) if pr.created_at else None,
    }


def _sse_line(evt: Dict[str, Any]) -> bytes:
    return ("data: " + json.dumps(evt, ensure_ascii=False) + "\n\n").encode("utf-8")


def _mock_stream(req: ParseReq, db: Session):
    structured = _mock_result(req.requirement)
    pieces = [
        "正在解析需求……\n",
        "提取源端信息 (SFTP)……\n",
        "提取目标端信息 (Hive)……\n",
        "提取任务基础信息……\n",
        "提取表结构字段……\n",
    ]
    for p in pieces:
        yield _sse_line({"type": "chunk", "content": p})
    content = json.dumps(structured, ensure_ascii=False)
    yield _sse_line({"type": "final", "final": True, "content": content, "structured": structured})
    pr = _persist_parse_result(
        db,
        req=req,
        content=content,
        structured=structured,
        raw_response=content,
        conv_id="mock-conv",
        msg_id="mock-msg",
        latency_ms=0,
        status="succeeded",
    )
    _audit(
        db,
        task_id=req.task_id,
        tool_name="unicom_ai_parse_collection_requirement.stream",
        request={"requirement": req.requirement, "mock": True},
        response={"structured": structured},
        status="succeeded",
    )
    db.commit()
    yield _sse_line({"type": "done", "parse_id": pr.id, "missing_fields": pr.missing_fields})


def _set_by_path(obj: Dict[str, Any], path: str, value: Any) -> None:
    parts = path.split(".")
    cur = obj
    for p in parts[:-1]:
        nxt = cur.get(p)
        if not isinstance(nxt, dict):
            nxt = {}
            cur[p] = nxt
        cur = nxt
    cur[parts[-1]] = value


@router.post("/docx")
async def extract_docx_plain_text(file: UploadFile = File(...)):
    """从 Word（.docx）提取纯文本，供采集需求录入回填；不做结构化解析。"""
    import io

    from docx import Document

    if not file.filename or not file.filename.lower().endswith(".docx"):
        raise HTTPException(400, "仅支持 .docx 文档")
    raw = await file.read()
    if not raw:
        raise HTTPException(400, "文件为空")
    try:
        doc = Document(io.BytesIO(raw))
    except Exception as e:
        raise HTTPException(400, f"无法读取 Word: {e}") from e
    lines: List[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            lines.append(t)
    for tb in doc.tables:
        for row in tb.rows:
            for cell in row.cells:
                t = (cell.text or "").strip()
                if t:
                    lines.append(t)
    text = "\n".join(lines)
    return {
        "text": text,
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
    }
